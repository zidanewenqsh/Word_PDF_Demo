from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_template():
    # 创建一个新的文档
    doc = Document()

    # 添加文档状态部分
    doc.add_heading('文档状态', 0)
    doc.add_paragraph('[ √ ]草稿')
    doc.add_paragraph('[  ]修订')
    doc.add_paragraph('[  ]发布')

    # 添加文档编号、管理部门等信息
    doc.add_paragraph('文档编号')
    doc.add_paragraph('管理部门')
    doc.add_paragraph('修订年月')
    doc.add_paragraph('版本号')

    # 添加项目可行性分析标题
    doc.add_heading('XXX项目 可行性分析', 0)

    # 添加修订人签字、审核人签字、批准人签字部分
    doc.add_paragraph('修订人签字')
    doc.add_paragraph('审核人签字')
    doc.add_paragraph('批准人签字')

    # 添加日期部分
    doc.add_paragraph('日期：')
    doc.add_paragraph('日期：')
    doc.add_paragraph('日期：')

    # 添加变更履历表格
    doc.add_heading('变更履历', 1)
    table = doc.add_table(rows=2, cols=8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '变更日期'
    hdr_cells[2].text = '版本'
    hdr_cells[3].text = '变更位置'
    hdr_cells[4].text = '变更原因'
    hdr_cells[5].text = '修订人'
    hdr_cells[6].text = '审核人'
    hdr_cells[7].text = '批准人'

    # 添加说明部分
    doc.add_paragraph('说明：“变更原因”主要是分为：')
    doc.add_paragraph('建立初稿')
    doc.add_paragraph('内容修订')
    doc.add_paragraph('正式发布')

    # 添加目录
    doc.add_heading('目录', 1)
    doc.add_paragraph('1. 文档介绍    4')
    doc.add_paragraph('   1.1. 编写目的    4')
    doc.add_paragraph('   1.2. 文档范围    4')
    doc.add_paragraph('   1.3. 读者对象    4')
    doc.add_paragraph('   1.4. 术语与缩写解释    4')
    doc.add_paragraph('   1.5. 参考资料    4')

    # 添加文档介绍部分
    doc.add_heading('文档介绍', 1)
    doc.add_paragraph('1.1. 编写目的')
    doc.add_paragraph('说明文档的编写目的')
    doc.add_paragraph('1.2. 文档范围')
    doc.add_paragraph('说明文档的主要内容')
    doc.add_paragraph('1.3. 读者对象')
    doc.add_paragraph('说明文档的读者对象')
    doc.add_paragraph('1.4. 术语与缩写解释')
    doc.add_paragraph('术语或缩写')
    doc.add_paragraph('解释')

    # 添加参考资料部分
    doc.add_heading('参考资料', 1)
    table = doc.add_table(rows=3, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '文档名称'
    hdr_cells[2].text = '文档编号'
    hdr_cells[3].text = '版本'
    hdr_cells[4].text = '发布日期'

    # 添加项目介绍部分
    doc.add_heading('项目介绍', 1)
    doc.add_paragraph('项目说明')
    doc.add_paragraph('介绍产品的名称、任务提出者、开发者、用户群')
    doc.add_paragraph('项目背景')
    doc.add_paragraph('介绍产品的背景，在什么样的背景下产生该产品')
    doc.add_paragraph('项目目标')
    doc.add_paragraph('介绍产品的目标与愿景，产品要能满足什么样的需求，要达到什么样一个效果')
    doc.add_paragraph('项目范围')
    doc.add_paragraph('略')
    doc.add_paragraph('项目风险')
    doc.add_paragraph('介绍产品的实施过程中，存在的内外部风险，如需求风险、技术风险等')
    doc.add_paragraph('假设与限制')
    doc.add_paragraph('介绍产品的外部环境因素条件与限制')

    # 添加现状与建设要求部分
    doc.add_heading('现状与建设要求', 1)
    doc.add_paragraph('用户分析')
    doc.add_paragraph('略')
    doc.add_paragraph('用户结构与管理职能')
    doc.add_paragraph('略')
    doc.add_paragraph('用户类别与其说明')
    table = doc.add_table(rows=3, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '用户类别'
    hdr_cells[2].text = '用户说明'
    hdr_cells[3].text = '角色简称'
    hdr_cells[4].text = '备注'

    # 添加用户环境部分
    doc.add_heading('用户环境', 1)
    doc.add_paragraph('操作系统')
    doc.add_paragraph('略')
    doc.add_paragraph('网络环境')
    doc.add_paragraph('略')
    doc.add_paragraph('与其他软件的兼容性')
    doc.add_paragraph('略')

    # 添加当前业务流程和数据状况部分
    doc.add_paragraph('当前业务流程')
    doc.add_paragraph('略')
    doc.add_paragraph('当前数据状况')
    doc.add_paragraph('略')

    # 添加建设任务与要求部分
    doc.add_paragraph('建设任务与要求')
    doc.add_paragraph('略')

    # 添加竞品项目现状部分
    doc.add_heading('竞品项目现状', 1)
    doc.add_paragraph('产品XXX1')
    doc.add_paragraph('运营公司XXX1')
    doc.add_paragraph('描述运营公司主体概要')
    doc.add_paragraph('业务开展')
    doc.add_paragraph('描述业务开展情况')
    doc.add_paragraph('公司网站')
    doc.add_paragraph('https://www.xxx1.com')
    doc.add_paragraph('产品XXX2')
    doc.add_paragraph('运营公司XXX2')
    doc.add_paragraph('描述运营公司主体概要')
    doc.add_paragraph('业务开展')
    doc.add_paragraph('描述业务开展情况')
    doc.add_paragraph('公司网站')
    doc.add_paragraph('https://www.xxx2.com')

    # 添加版权声明
    doc.add_paragraph('北京中盈安信技术服务有限公司版权所有')
    doc.add_paragraph('本文档中所包含的信息属于北京中盈安信技术服务有限公司的机密信息')
    doc.add_paragraph('未经许可，不可全部或部分发表、复制、使用于任何目的')
    doc.add_paragraph('XXX有限公司版权所有')
    doc.add_paragraph('本文档中所包含的信息属于XXX有限公司的机密信息')
    doc.add_paragraph('未经许可，不可全部或部分发表、复制、使用于任何目的')

    # 保存文档
    doc.save('可行性分析 - 模板.docx')

create_template()