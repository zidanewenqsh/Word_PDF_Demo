from docx import Document
# 读取 Word 文件中的所有段落
# 打开并读取 Word 文件
doc = Document('example3.docx')

# 获取文档中的所有段落
for para in doc.paragraphs:
    print(para.text)
    if para.style.name.startswith('Heading'):  # 识别标题
        print(f'{para.style.name} 标题: {para.text}')
    for run in para.runs:  # 检查每个run的格式
        print(f'  Run Text: {run.text}, Bold: {run.bold}, Italic: {run.italic}, Font Size: {run.font.size}')

# 读取 Word 文件中的表格
# 获取文档中的所有表格
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text, end=' | ')
        # print()  # 换行

# 获取第一个表格
table = doc.tables[0]

# 读取表格的每一行和每个单元格
for row in table.rows:
    row_data = [cell.text.strip() for cell in row.cells]
    print(row_data)


# 获取所有段落，检查是否为无序列表
for para in doc.paragraphs:
    # print(para.style.name)
    if para.style.name == 'List Bullet':  # 无序列表
        print(f'无序列表项: {para.text}')

# 获取所有段落，检查是否为有序列表
for para in doc.paragraphs:
    if para.style.name == 'List Number':  # 有序列表
        print(f'有序列表项: {para.text}')

# 获取文档中的所有图片
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:  # 判断是否为图片文件
        print(f'图片路径: {rel.target_ref}')

# 获取页眉和页脚
section = doc.sections[0]

# 读取页眉内容
header = section.header
for para in header.paragraphs:
    print(f'页眉: {para.text}')

# 读取页脚内容
footer = section.footer
for para in footer.paragraphs:
    print(f'页脚: {para.text}')

doc.save('example2.docx')  # 保存文档
