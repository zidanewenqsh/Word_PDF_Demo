from docx import Document
from docx.shared import Pt, RGBColor
from matplotlib import mathtext
import matplotlib.pyplot as plt

doc = Document()

# 使用matplotlib绘制公式
fig, ax = plt.subplots()
ax.text(0.5, 0.5, r'$\frac{a}{b} = c$', fontsize=30, ha='center')
ax.axis('off')  # 不显示坐标轴
fig.savefig('equation.png')

# 创建Word文档并插入公式图片
doc = Document()
doc.add_paragraph("这是一个插入的公式：")
doc.add_picture('equation.png')


doc.save('example4.docx')  # 保存文档
# doc.close()  # 关闭文档

