from docx import Document
from docx.shared import Pt, RGBColor
# 读取 Word 文件中的所有段落
# 打开并读取 Word 文件
doc = Document()

# 添加一个加粗的段落
para = doc.add_paragraph("这是一个加粗的段落。")
para.bold = True

# 添加一个斜体且字体大小为14号的段落
para2 = doc.add_paragraph("这是一个斜体且字体为14号的段落。")
para2.italic = True
para2.runs[0].font.size = Pt(14)

# 设置字体颜色为红色
para3 = doc.add_paragraph("这是一个字体颜色为红色的段落。")
para3.runs[0].font.color.rgb = RGBColor(255, 0, 0)

# 添加一级标题
doc.add_heading('这是一级标题', level=1)

# 添加二级标题
doc.add_heading('这是二级标题', level=2)

# 添加三级标题
doc.add_heading('这是三级标题', level=3)

# 添加四级标题
doc.add_heading('这是四级标题', level=4)

# 添加无序列表（点符号列表）
doc.add_paragraph('无序列表项 1', style='List Bullet')
doc.add_paragraph('无序列表项 2', style='List Bullet')
doc.add_paragraph('无序列表项 3', style='List Bullet')

# 添加有序列表（编号列表）
doc.add_paragraph('有序列表项 1', style='List Number')
doc.add_paragraph('有序列表项 2', style='List Number')
doc.add_paragraph('有序列表项 3', style='List Number')

# 插入图片
doc.add_picture('resources/mycat.jpg', width=Pt(100))

# 获取页眉
section = doc.sections[0]
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = '这是页眉内容'

# 获取页脚
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = '这是页脚内容'

# 添加一个3x3的表格
table = doc.add_table(rows=3, cols=3)

# 填充表格内容并设置单元格背景色
table.cell(0, 0).text = "Header 1"
table.cell(0, 1).text = "Header 2"
table.cell(0, 2).text = "Header 3"

# 为其他单元格添加内容
table.cell(1, 0).text = "Cell 1,1"
table.cell(1, 1).text = "Cell 1,2"
table.cell(1, 2).text = "Cell 1,3"

# 设置第一行的背景颜色
table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True


doc.save('example3.docx')  # 保存文档
# doc.close()  # 关闭文档